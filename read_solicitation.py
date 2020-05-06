
from docx import Document
import PyPDF2
from flask import jsonify
import sys
import re


# if __name__ == '__main__':
# def read_solicitation(file):
#     if file.endswith(".docx"):
#
#         doc = Document(file)
#
#         found_sow = False
#         extract = list()
#         for p in doc.paragraphs:
#             #print('-------')
#     #         print(p.text)
#             if p.text.lower().endswith('statement of work') or p.text.lower().endswith('statement of work (sow)'):
#                 found_sow = True
#             elif found_sow:
#                 if p.text.lower().startswith('title:') or p.text.lower().startswith('project title:'):
#                     extract.append(p.text)
#                 elif p.text.lower().startswith('background:'):
#                     extract.append(p.text)
#                 elif p.text.lower().startswith('scope:'):
#                     extract.append(p.text)
#             result = '\n\n'.join(extract)
#         print(result)
#         return((result))
#         return(result)
#
#     elif file.endswith(".pdf"):
#         object = PyPDF2.PdfFileReader(file)
#         NumPages = object.getNumPages()
#         # extract text and do the search
#         for i in range(0, NumPages):
#             PageObj = object.getPage(i)
#             Text = PageObj.extractText().lower().replace('\n', '')
#             ResSearch = re.findall(r"(.{0,100}%s.{0,100})" % 'housing', Text, re.MULTILINE)
#             if ResSearch == []:
#                 continue
#             result = print("Occurence(s) on page " + str(i) + ":")
#             for each in ResSearch:
#         #         print(ResSearch)
#                 for elem in ResSearch:
#                     result += print('...', elem, '...')
#                     result += print('\n')
#             print(result)
#             return(result)




def read_solicitation(file):
    doc = Document(file)
    found_sow = False
    extract = list()
    for p in doc.paragraphs:
        if p.text.lower().endswith('statement of work') or p.text.lower().endswith('statement of work (sow)'):
            found_sow = True
        elif found_sow:
            if p.text.lower().startswith('title:') or p.text.lower().startswith('project title:'):
                extract.append(p.text)
            elif p.text.lower().startswith('background:'):
                extract.append(p.text)
            elif p.text.lower().startswith('scope:'):
                extract.append(p.text)
        result = '\n\n'.join(extract)
    print(result)
    return((result))
    return(result)






read_solicitation('Sample_SOW.docx')
